#!/usr/bin/env python3
"""
PPT Master - Planning Document Renderer

Renders `plan_spec.md` into the 기획서 a person actually reads — Markdown
always, Word (.docx) on request. The frame's section chain is kept exactly as
written: this script formats the contract, it never authors, reorders or
rewrites it.

`plan_spec.md` is the source of truth for both halves of the run. The deck
takes that material through `outline.py` and the golden circle; this is the
other half — the report that falls out of the same planning pass, in the
frame's own order.

Unfinished work stays visible. A section at `초안` or `확인 필요` is rendered
with that state on it rather than silently reading as settled, and the sections
still open are listed at the end.

Usage:
    python3 scripts/render_plan_doc.py <project_path> [--format auto|md|docx|both]

Examples:
    python3 scripts/render_plan_doc.py projects/20260903_speakup_campaign
    python3 scripts/render_plan_doc.py projects/20260903_speakup_campaign --format both
    python3 scripts/render_plan_doc.py projects/20260903_speakup_campaign --format docx

Dependencies:
    python-docx  - only for --format docx / both. Imported lazily, so the
                   default Markdown path runs on the standard library alone.

See references/planner.md §6 for the contract this renders, and
references/report-format.md for the Word form and its register.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).resolve().parent))

from console_encoding import configure_utf8_stdio  # noqa: E402
from plan_spec import FRAMES, Frame, Section, parse  # noqa: E402


DOC_SUFFIX = "기획서"

# Report kinds that ask for a document, not just a deck (intake.doc_kind).
DOC_KINDS_WANTING_WORD = {"보고서", "둘 다"}

TITLE_FONT = "맑은 고딕"   # 제목·대제목 (report-format.md §4)
BODY_FONT = "한컴바탕"     # 본문·표·각주

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_BULLET_RE = re.compile(r"^\s*[-*·]\s+(.*)$")


@dataclass
class Plan:
    """Everything the renderers need, read once from the project."""

    title: str
    meta: dict
    sections: list[Section]
    frame: Optional[Frame]
    intake: dict


def load_intake(project: Path) -> dict:
    """Read `intake.json` if it is there. Rendering does not depend on it."""
    path = project / "intake.json"
    if not path.is_file():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        print(f"[render_plan_doc] intake.json unreadable ({exc}) — "
              f"rendering without it", file=sys.stderr)
        return {}
    return data if isinstance(data, dict) else {}


def build_plan(project: Path) -> Plan:
    """Parse `plan_spec.md` into the shape both renderers consume."""
    text = (project / "plan_spec.md").read_text(encoding="utf-8")
    meta, sections = parse(text)
    heading = re.search(r"^#\s+(.+)$", text, re.M)
    frame = FRAMES.get(str(meta.get("frame", "")).strip())
    if frame is None and meta.get("frame"):
        print(f"[render_plan_doc] unknown frame {meta.get('frame')!r} in "
              f"plan_spec.md — rendering the sections as written", file=sys.stderr)
    return Plan(
        title=heading.group(1).strip() if heading else DOC_SUFFIX,
        meta=meta,
        sections=sections,
        frame=frame,
        intake=load_intake(project),
    )


def header_rows(plan: Plan) -> list[tuple[str, str]]:
    """The identifying block at the top of the document."""
    rows: list[tuple[str, str]] = []
    if plan.frame is not None:
        rows.append(("구성", f"{plan.frame.label} · {len(plan.frame.sections)}단"))
    for label, key in (("목적", "purpose"), ("출발", "assignment")):
        value = str(plan.meta.get(key) or plan.intake.get(key) or "").strip()
        if value:
            rows.append((label, value))
    audience = str(plan.intake.get("audience") or "").strip()
    if audience:
        rows.append(("대상", audience))
    rows.append(("작성", _stamp_date(str(plan.meta.get("generated_at") or ""))))
    return rows


def _stamp_date(stamp: str) -> str:
    """`generated_at` as a plain date; today's when the stamp is unreadable."""
    try:
        return datetime.fromisoformat(stamp).strftime("%Y-%m-%d")
    except ValueError:
        return datetime.now().strftime("%Y-%m-%d")


def _empty_note(plan: Plan, section: Section) -> str:
    """Why a section is blank — the frame's own reason, not a guess."""
    if plan.frame is not None and section.name in plan.frame.fact_required:
        return "자료 근거 없음 — 미작성"
    return "미작성"


def pending_sections(plan: Plan) -> list[Section]:
    return [s for s in plan.sections if s.status != "확정"]


# ---- Markdown ---------------------------------------------------------------


def render_markdown(plan: Plan) -> str:
    out: list[str] = [f"# {plan.title}", ""]

    rows = header_rows(plan)
    if rows:
        out += ["| | |", "|---|---|"]
        out += [f"| {k} | {v} |" for k, v in rows]
        out.append("")

    for i, sec in enumerate(plan.sections, 1):
        out += [f"## {i}. {sec.name}", ""]
        if sec.status != "확정":
            out += [f"> **{sec.status}**", ""]
        if sec.body:
            out += [sec.body, ""]
        else:
            out += [f"_{_empty_note(plan, sec)}_", ""]
        if sec.source:
            out += [f"근거: `{sec.source}`", ""]

    pending = pending_sections(plan)
    if pending:
        out += ["---", "", "## 아직 닫히지 않은 항목", ""]
        out += [f"- {s.name} — {s.status}" for s in pending]
        out.append("")

    return "\n".join(out).rstrip() + "\n"


# ---- Word — the report form (references/report-format.md) -------------------

# §3 marker hierarchy. Sections are the 대제목 level, so they take the 번호박스.
ROMAN = "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ"
M_CORE = "\u25a1"      # □ 핵심 — the 두괄식 opening line
M_DETAIL = "\u2500"    # ─ 세부
M_NOTE = "*"            # 각주

# §5. Figures carry weight, never hue: 개선/악화 cannot be read off a number
# without knowing the indicator's direction, and guessing reverses the meaning.
_NUMBER_RE = re.compile(
    r"\d[\d,.]*\s*(?:%p|%|건|명|원|배|점|일|주|개월|년|분기|억|만원|천원)?"
)


def render_docx(plan: Plan, out_path: Path) -> None:
    """Write the 기획서 in the report form. Raises RuntimeError without python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed — it is needed only for --format docx/both. "
            "Install it with:  pip install python-docx"
        ) from exc

    doc = Document()
    _setup_page(doc, Mm)
    _base_style(doc, qn, Pt)

    # 제목 — 맑은 고딕 17pt, centred (§4).
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_after = Pt(14)
    _font(head.add_run(plan.title), qn, Pt, name=TITLE_FONT, size=17, bold=True)

    _meta_table(doc, plan, qn, Pt)

    for i, sec in enumerate(plan.sections):
        _section_heading(doc, i, sec, qn, Pt, OxmlElement)

        blocks = body_blocks(sec.body) if sec.body else []
        if blocks:
            # First block is the 두괄식 conclusion; an option line is its own
            # 핵심; everything else supports them (§6).
            for n, (text, _is_bullet) in enumerate(blocks):
                core = n == 0 or text.startswith("**")
                _marked(doc, M_CORE if core else M_DETAIL, text, qn, Pt,
                        size=11, indent=Mm(0 if core else 4))
        else:
            _marked(doc, M_DETAIL, _empty_note(plan, sec), qn, Pt,
                    size=11, indent=Mm(4))

        if sec.source:
            _marked(doc, M_NOTE, f" 근거: {sec.source}", qn, Pt, size=9,
                    indent=Mm(4), plain=True)

    pending = pending_sections(plan)
    if pending:
        block = doc.add_paragraph()
        block.paragraph_format.space_before = Pt(16)
        block.paragraph_format.space_after = Pt(4)
        _font(block.add_run("< 아직 닫히지 않은 항목 >"), qn, Pt,
              name=TITLE_FONT, size=12, bold=True)
        for sec in pending:
            _marked(doc, M_DETAIL, f"{sec.name} — {sec.status}", qn, Pt,
                    size=11, indent=Mm(4))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def body_blocks(body: str) -> list[tuple[str, bool]]:
    """Regroup a section body into paragraphs, with a bullet flag.

    A single newline inside `plan_spec.md` is a soft wrap, so joining those
    lines keeps one sentence one paragraph. Bullets and option labels
    (`**1안** …`) own their line by contract and stay separate.
    """
    blocks: list[tuple[str, bool]] = []
    buffered: list[str] = []

    def flush() -> None:
        if buffered:
            blocks.append((" ".join(buffered), False))
            buffered.clear()

    for line in body.split("\n"):
        stripped = line.strip()
        bullet = _BULLET_RE.match(line)
        if not stripped:
            flush()
        elif bullet:
            flush()
            blocks.append((bullet.group(1).strip(), True))
        elif stripped.startswith("**"):
            flush()
            blocks.append((stripped, False))
        else:
            buffered.append(stripped)
    flush()
    return blocks


def _setup_page(doc, Mm) -> None:
    """A4 portrait (§4)."""
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)


def _base_style(doc, qn, Pt) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), BODY_FONT)
    style.paragraph_format.space_after = Pt(2)


def _font(run, qn, Pt, *, name: str, size: float, bold: bool = False):
    """Name the face for Latin and East Asian text — Korean falls back otherwise."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    return run


def _meta_table(doc, plan: Plan, qn, Pt) -> None:
    """The identifying block as a 통계표 — 9pt, ruled (§4)."""
    rows = header_rows(plan)
    if not rows:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        for cell, text, bold in ((cells[0], key, True), (cells[1], value, False)):
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            _font(para.add_run(text), qn, Pt, name=BODY_FONT, size=9, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _section_heading(doc, index: int, section: Section, qn, Pt, OxmlElement) -> None:
    """`Ⅰ` in a box, then the section name, then its state when it is not 확정."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(5)

    numeral = ROMAN[index] if index < len(ROMAN) else str(index + 1)
    box = _font(para.add_run(f" {numeral} "), qn, Pt, name=TITLE_FONT, size=12, bold=True)
    _boxed(box, qn, OxmlElement)

    _font(para.add_run(f"  {section.name}"), qn, Pt, name=TITLE_FONT, size=12, bold=True)
    if section.status != "확정":
        _font(para.add_run(f"  ({section.status})"), qn, Pt, name=BODY_FONT, size=9)


def _boxed(run, qn, OxmlElement) -> None:
    """Rule a box around one run — the 번호박스 of the 대제목."""
    border = OxmlElement("w:bdr")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "2")
    border.set(qn("w:color"), "000000")
    run._element.get_or_add_rPr().append(border)


def _marked(doc, marker: str, text: str, qn, Pt, *, size: float, indent,
            plain: bool = False) -> None:
    """One marker line: the marker, then the text with its figures in bold.

    `plain` writes the text untouched — a 각주 carries a path and a line range,
    and emphasising the digits inside `L18-L44` reads them as findings.
    """
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = indent
    para.paragraph_format.space_after = Pt(2)
    _font(para.add_run(f"{marker} "), qn, Pt, name=BODY_FONT, size=size)
    spans = [(text, False)] if plain else _spans(text)
    for chunk, bold in spans:
        _font(para.add_run(chunk), qn, Pt, name=BODY_FONT, size=size, bold=bold)


def _spans(text: str) -> list[tuple[str, bool]]:
    """Split into runs, bolding `**marked**` spans and every figure (§5)."""
    out: list[tuple[str, bool]] = []
    cursor = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > cursor:
            out += _figure_spans(text[cursor:match.start()])
        out.append((match.group(1), True))
        cursor = match.end()
    out += _figure_spans(text[cursor:])
    return [(chunk, bold) for chunk, bold in out if chunk]


def _figure_spans(text: str) -> list[tuple[str, bool]]:
    out: list[tuple[str, bool]] = []
    cursor = 0
    for match in _NUMBER_RE.finditer(text):
        if match.start() > cursor:
            out.append((text[cursor:match.start()], False))
        out.append((match.group(0), True))
        cursor = match.end()
    out.append((text[cursor:], False))
    return out


# ---- CLI --------------------------------------------------------------------


def resolve_format(requested: str, intake: dict) -> str:
    """`auto` follows what the person said they wanted at intake."""
    if requested != "auto":
        return requested
    kind = str(intake.get("doc_kind") or "").strip()
    return "both" if kind in DOC_KINDS_WANTING_WORD else "md"


def main(argv: Optional[list[str]] = None) -> int:
    configure_utf8_stdio()
    args = build_parser().parse_args(argv)

    project = Path(args.project_path).resolve()
    if not project.is_dir():
        print(f"[render_plan_doc] project not found: {project}", file=sys.stderr)
        return 1

    plan_path = project / "plan_spec.md"
    if not plan_path.is_file():
        print(f"[render_plan_doc] plan_spec.md not found at {plan_path} — "
              f"run plan_spec.py --scaffold first", file=sys.stderr)
        return 1

    plan = build_plan(project)
    fmt = resolve_format(args.format, plan.intake)
    out_dir = project / "exports"
    stem = f"{project.name}_{DOC_SUFFIX}"
    written: list[Path] = []

    if fmt in {"md", "both"}:
        out_dir.mkdir(parents=True, exist_ok=True)
        md_path = out_dir / f"{stem}.md"
        md_path.write_text(render_markdown(plan), encoding="utf-8")
        written.append(md_path)

    if fmt in {"docx", "both"}:
        docx_path = out_dir / f"{stem}.docx"
        try:
            render_docx(plan, docx_path)
            written.append(docx_path)
        except RuntimeError as exc:
            # Asked for explicitly: refuse rather than quietly hand back less.
            # Chosen by `auto`: the Markdown is already written, so say what is
            # missing and how to get it instead of failing the run.
            if args.format != "auto":
                print(f"[render_plan_doc] {exc}", file=sys.stderr)
                return 1
            print(f"[render_plan_doc] Word skipped — {exc}", file=sys.stderr)

    pending = pending_sections(plan)
    label = plan.frame.label if plan.frame else "알 수 없는 구성"
    print(f"[render_plan_doc] {label} · {len(plan.sections)}절 · "
          f"닫히지 않은 절 {len(pending)}개", file=sys.stderr)
    for path in written:
        print(str(path))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Render plan_spec.md into the 기획서 document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("project_path", help="Path to project directory")
    parser.add_argument(
        "--format", choices=("auto", "md", "docx", "both"), default="auto",
        help="auto follows intake.doc_kind (보고서/둘 다 → both, otherwise md)",
    )
    return parser


if __name__ == "__main__":
    raise SystemExit(main())
