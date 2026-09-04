#!/usr/bin/env python3
"""
PPT Master - Report Form Writer

Typesets a report into Word in one of the forms under
`templates/report_forms/`. The form supplies the marker glyphs, fonts, sizes
and colours; this module supplies the mechanics — marker hierarchy, tables,
badges, page chrome.

It formats what it is handed and authors nothing. Section bodies arrive already
written in the 개조식 register with their markers in the text; parsing them back
out is typesetting, not rewriting.

Usage:
    Imported by scripts/render_plan_doc.py. Not a CLI.

Dependencies:
    python-docx (the caller imports it lazily and reports the install line)

See references/report-format.md for the form contract this implements.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Optional

if __name__ == "__main__":
    print(__doc__)
    raise SystemExit(0)

FORMS_DIR = Path(__file__).resolve().parent.parent / "templates" / "report_forms"

ROMAN = "ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ"

# Body markers the planner writes (report-format.md §4). Both the 한국은행 and
# the 한수원 glyphs are accepted whichever form is rendering — the form decides
# what is drawn, the author should not have to know which.
_CORE_RE = re.compile(r"^\s*[▢□■]\s*(.*)$")
_DETAIL_RE = re.compile(r"^\s*[◦○─―∙•\-\*]\s+(.*)$")
_NOTE_RE = re.compile(r"^\s*(?:\*(?!\*)|※)\s*(.*)$")
_TABLE_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|[\s:\-|]+\|\s*$")

# Inline markup. `{+…}` / `{-…}` carry the direction the author asserts —
# 개선 or 악화 — because a figure alone cannot say which (report-format.md §5).
_INLINE_RE = re.compile(
    r"\{([+-])([^}]*)\}"                      # direction
    r"|\*\*(.+?)\*\*"                         # bold
    r"|\[([^\]]*(?:확인 필요|추정)[^\]]*)\]"   # badge
)
# A figure is the number *and* the unit it is stated in — "18.50건" is one thing,
# and bolding half of it splits the value from what it measures.
_UNITS = (r"%p|%|건|명|원|점|개월|년|월|일|주|시간|분|초|회|차|개|배|억원|억"
          r"|만원|천원|만|천|배수|명당|건당")
_ARROW = r"(?:[▲▼△▽]\s?)?"   # a space belongs to the arrow, never to a bare figure
# In prose a bare numeral is usually part of a name (1공구, 문항 3) — only a
# number carrying a unit is a figure. In a table the number *is* the value.
_FIGURE_RE = re.compile(_ARROW + r"\d[\d,.]*\s*(?:" + _UNITS + r")")
_FIGURE_ANY_RE = re.compile(_ARROW + r"\d[\d,.]*\s*(?:" + _UNITS + r")?")


def load_form(name: str) -> dict:
    """Read one form spec. `auto` resolves through the index's default."""
    index = json.loads((FORMS_DIR / "forms_index.json").read_text(encoding="utf-8"))
    key = index.get("default", "bok") if name in ("", "auto") else name
    path = FORMS_DIR / f"{key}.json"
    if not path.is_file():
        known = ", ".join(sorted(index.get("forms", {})))
        raise ValueError(f"unknown report form {key!r} — known forms: {known}")
    return json.loads(path.read_text(encoding="utf-8"))


# ---- body → blocks ----------------------------------------------------------


def parse_body(body: str) -> list[tuple[str, Any]]:
    """Split a section body into ('core'|'detail'|'note'|'table', payload).

    A body written without markers keeps the old shape: the opening paragraph
    is the 두괄식 핵심 and the rest supports it.
    """
    blocks: list[tuple[str, Any]] = []
    lines = body.split("\n")
    marked = any(_CORE_RE.match(ln) or _NOTE_RE.match(ln) for ln in lines)
    buffered: list[str] = []
    table: list[list[str]] = []

    def flush_text() -> None:
        if buffered:
            text = " ".join(buffered)
            level = "core" if (not marked and not blocks) else "detail"
            blocks.append((level, text))
            buffered.clear()

    def flush_table() -> None:
        if table:
            blocks.append(("table", [row[:] for row in table]))
            table.clear()

    for line in lines:
        stripped = line.strip()
        if _TABLE_RE.match(line):
            flush_text()
            if not _TABLE_SEP_RE.match(line):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                table.append(cells)
            continue
        flush_table()

        if not stripped:
            flush_text()
            continue
        if (m := _CORE_RE.match(line)):
            flush_text()
            blocks.append(("core", m.group(1).strip()))
        elif stripped.startswith("**"):
            # An option label (`**1안** …`) leads its own 핵심. Tested before the
            # 각주 rule, which would otherwise claim the first `*` of the pair.
            flush_text()
            blocks.append(("core", stripped))
        elif (m := _NOTE_RE.match(line)):
            flush_text()
            blocks.append(("note", m.group(1).strip()))
        elif (m := _DETAIL_RE.match(line)):
            flush_text()
            blocks.append(("detail", m.group(1).strip()))
        else:
            buffered.append(stripped)

    flush_text()
    flush_table()
    return blocks


def strip_markup(text: str) -> str:
    """Plain text for the Markdown draft — braces are a Word-only instruction."""
    return _INLINE_RE.sub(
        lambda m: m.group(2) or m.group(3) or (f"[{m.group(4)}]" if m.group(4) else ""),
        text,
    )


def _runs(text: str, *, figures: bool = True,
          bare: bool = False) -> list[tuple[str, str]]:
    """Text split into (chunk, role) — '', 'bold', 'improve', 'worsen', 'badge'.

    `figures` off writes the text untouched — a 각주 carries a path and a line
    range, and emphasising the digits inside `L12-L88` reads them as findings.
    `bare` on also bolds a unit-less number, which is what a table cell holds.
    """
    out: list[tuple[str, str]] = []
    cursor = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > cursor:
            out += _plain(text[cursor:m.start()], figures, bare)
        if m.group(1):
            out.append((m.group(2), "improve" if m.group(1) == "+" else "worsen"))
        elif m.group(3):
            out.append((m.group(3), "bold"))
        else:
            out.append((m.group(4), "badge"))
        cursor = m.end()
    out += _plain(text[cursor:], figures, bare)
    return [(chunk, role) for chunk, role in out if chunk]


def _plain(text: str, figures: bool, bare: bool = False) -> list[tuple[str, str]]:
    if not figures:
        return [(text, "")]
    out: list[tuple[str, str]] = []
    cursor = 0
    for m in (_FIGURE_ANY_RE if bare else _FIGURE_RE).finditer(text):
        if m.start() > cursor:
            out.append((text[cursor:m.start()], ""))
        out.append((m.group(0), "bold"))
        cursor = m.end()
    out.append((text[cursor:], ""))
    return out


# ---- Word -------------------------------------------------------------------


class ReportWriter:
    """Holds the docx handles so every helper does not re-import them."""

    def __init__(self, form: dict):
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Emu, Mm, Pt, RGBColor

        self.form = form
        self.doc = Document()
        self.Mm, self.Pt, self.RGBColor, self.Emu = Mm, Pt, RGBColor, Emu
        self.qn, self.OxmlElement = qn, OxmlElement
        self.ALIGN, self.TABLE_ALIGN = WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT
        self.TAB_RIGHT = WD_TAB_ALIGNMENT.RIGHT
        self._setup_page()

    # -- primitives --

    def rgb(self, key: str):
        return self.RGBColor.from_string(self.form["colors"][key])

    def font(self, run, *, name: str, size: float, bold: bool = False,
             color: Optional[str] = None):
        run.font.name = name
        run.font.size = self.Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = self.RGBColor.from_string(color)
        rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(self.qn(attr), name)
        return run

    def shade(self, element, fill: str) -> None:
        shd = self.OxmlElement("w:shd")
        shd.set(self.qn("w:val"), "clear")
        shd.set(self.qn("w:fill"), fill)
        element.append(shd)

    def bottom_rule(self, para, color: str, size: int = 12) -> None:
        borders = self.OxmlElement("w:pBdr")
        bottom = self.OxmlElement("w:bottom")
        bottom.set(self.qn("w:val"), "single")
        bottom.set(self.qn("w:sz"), str(size))
        bottom.set(self.qn("w:space"), "3")
        bottom.set(self.qn("w:color"), color)
        borders.append(bottom)
        para._p.get_or_add_pPr().append(borders)

    def _no_auto_space(self) -> None:
        """Stop Word padding the seam between Hangul and Latin.

        On by default, it renders "1,015명" as "1,015 명" — the figure split from
        the unit it is stated in. Set on the document defaults, where `w:pPr` is
        empty and the schema's element order cannot be violated.
        """
        styles = self.doc.styles.element
        defaults = styles.find(self.qn("w:docDefaults"))
        if defaults is None:
            defaults = self.OxmlElement("w:docDefaults")
            styles.insert(0, defaults)
        ppr_default = defaults.find(self.qn("w:pPrDefault"))
        if ppr_default is None:
            ppr_default = self.OxmlElement("w:pPrDefault")
            defaults.append(ppr_default)
        ppr = ppr_default.find(self.qn("w:pPr"))
        if ppr is None:
            ppr = self.OxmlElement("w:pPr")
            ppr_default.append(ppr)
        # `w:pPr` is an ordered sequence and both elements precede `w:spacing`,
        # which python-docx has already written — so insert at the front.
        for tag in ("w:autoSpaceDN", "w:autoSpaceDE"):
            el = self.OxmlElement(tag)
            el.set(self.qn("w:val"), "0")
            ppr.insert(0, el)

    def no_auto_space_para(self, para) -> None:
        """Repeat the setting on one paragraph.

        The document default is what Word reads, and it is enough for Word. Other
        readers — LibreOffice among them — ignore `docDefaults` here and pad the
        seam anyway, which means a converted preview shows "1,015 명" and nobody
        can tell a real defect from a rendering artifact. Stating it per
        paragraph makes the file say the same thing to every reader.
        """
        ppr = para._p.get_or_add_pPr()
        for tag in ("w:autoSpaceDN", "w:autoSpaceDE"):
            if ppr.find(self.qn(tag)) is not None:
                continue
            el = self.OxmlElement(tag)
            el.set(self.qn("w:val"), "0")
            ppr.insert(0, el)

    def top_rule(self, para, color: str, size: int = 6) -> None:
        borders = self.OxmlElement("w:pBdr")
        top = self.OxmlElement("w:top")
        top.set(self.qn("w:val"), "single")
        top.set(self.qn("w:sz"), str(size))
        top.set(self.qn("w:space"), "6")
        top.set(self.qn("w:color"), color)
        borders.append(top)
        para._p.get_or_add_pPr().append(borders)

    def _setup_page(self) -> None:
        page = self.form["page"]
        s = self.doc.sections[0]
        s.page_width = self.Mm(page["width_mm"])
        s.page_height = self.Mm(page["height_mm"])
        s.top_margin = self.Mm(page["margin_top_mm"])
        s.bottom_margin = self.Mm(page["margin_bottom_mm"])
        s.left_margin = s.right_margin = self.Mm(page["margin_side_mm"])
        style = self.doc.styles["Normal"]
        style.font.name = self.form["fonts"]["body"]
        style.font.size = self.Pt(self.form["sizes"]["detail"])
        style.element.get_or_add_rPr().get_or_add_rFonts().set(
            self.qn("w:eastAsia"), self.form["fonts"]["body"])
        style.paragraph_format.space_after = self.Pt(1)
        style.paragraph_format.line_spacing = self.form["line_spacing"]["base"]
        self._no_auto_space()

    # -- content --

    def write_line(self, level: str, text: str, size: Optional[float] = None) -> None:
        f, sizes, colors = self.form, self.form["sizes"], self.form["colors"]
        marker = f["markers"].get(level, "")
        if size is None:
            size = sizes["core"] if level == "core" else \
                sizes["note"] if level == "note" else sizes["detail"]
        indent = {"core": 0.0, "detail": 4.0, "note": 8.0}[level]

        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = self.Mm(indent + 4)
        pf.first_line_indent = self.Mm(-4)
        pf.line_spacing = self.form["line_spacing"]["list"]
        pf.space_after = self.Pt(0)
        pf.space_before = self.Pt(4 if level == "core" else 0)

        body_color = colors["text"] if level != "core" else "1F2937"
        # A 핵심 line is bold as a whole — unless the author already said which
        # part carries the weight, in which case the markup wins.
        blanket_bold = level == "core" and "**" not in text
        self.font(para.add_run(f"{marker} "), name=f["fonts"]["body"], size=size,
                  bold=level == "core", color=body_color)
        for chunk, role in _runs(text, figures=not text.lstrip().startswith("근거:")):
            run = self.font(
                para.add_run(chunk), name=f["fonts"]["body"], size=size,
                bold=blanket_bold or role in ("bold", "improve", "worsen", "badge"),
                color=colors["improve"] if role == "improve"
                else colors["worsen"] if role == "worsen"
                else colors["badge_fg"] if role == "badge"
                else body_color,
            )
            if role == "badge":
                self.shade(run._element.get_or_add_rPr(), colors["badge_bg"])

    def write_table(self, rows: list[list[str]], size: Optional[float] = None) -> None:
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        if not rows:
            return
        cell_size = size or sizes["table"]
        width = max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=width)
        table.alignment = self.TABLE_ALIGN.CENTER
        table.autofit = False
        self._table_borders(table)

        for i, row in enumerate(rows):
            cells = table.add_row().cells
            for j in range(width):
                text = row[j] if j < len(row) else ""
                cell = cells[j]
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = self.Pt(0)
                para.paragraph_format.line_spacing = self.form["line_spacing"]["base"]
                para.alignment = self.ALIGN.LEFT if j == 0 else self.ALIGN.CENTER
                if i == 0:
                    self.shade(cell._tc.get_or_add_tcPr(), colors["table_header_bg"])
                    self.font(para.add_run(text), name=f["fonts"]["body"],
                              size=cell_size, bold=True,
                              color=colors["table_header_fg"])
                    para.alignment = self.ALIGN.CENTER
                    continue
                if j == 0:
                    self.shade(cell._tc.get_or_add_tcPr(), colors["table_label_bg"])
                for chunk, role in _runs(text, bare=True):
                    self.font(
                        para.add_run(chunk), name=f["fonts"]["body"],
                        size=cell_size,
                        bold=j == 0 or role in ("bold", "improve", "worsen"),
                        color=colors["improve"] if role == "improve"
                        else colors["worsen"] if role == "worsen"
                        else colors["text"],
                    )
        self._table_widths(table, width, rows)
        self.doc.add_paragraph().paragraph_format.space_after = self.Pt(2)

    def _table_widths(self, table, cols: int, rows: list) -> None:
        """Give the label column the room its text needs, split the rest evenly."""
        section = self.doc.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin
        # Share the width by what each column actually holds, so a long label
        # column does not wrap while a 증감 column sits half empty. Hangul is
        # about twice the advance of a Latin character.
        def weight(text: str) -> int:
            return sum(2 if ord(c) > 0x1100 else 1 for c in strip_markup(text))

        # Column widths are only honoured under a fixed layout; the default
        # autofit recomputes them from content and discards what we set.
        layout = self.OxmlElement("w:tblLayout")
        layout.set(self.qn("w:type"), "fixed")
        table._tbl.tblPr.append(layout)

        demand = [max((weight(row[j]) for row in rows if j < len(row)), default=1)
                  for j in range(cols)]
        floor = 0.10
        share = [max(d / sum(demand), floor) for d in demand]
        share = [v / sum(share) for v in share]
        widths = [self.Emu(int(usable * v)) for v in share]
        grid = table._tbl.find(self.qn("w:tblGrid"))
        if grid is not None:
            for j, col in enumerate(grid.findall(self.qn("w:gridCol"))):
                if j < len(widths):
                    col.set(self.qn("w:w"), str(int(widths[j].twips)))
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(widths):
                    cell.width = widths[j]

    def _table_borders(self, table) -> None:
        """Horizontal rules only unless the form keeps vertical ones."""
        colors = self.form["colors"]
        vertical = self.form["rules"].get("table_vertical_rules", False)
        borders = self.OxmlElement("w:tblBorders")
        spec = {
            "top": ("single", "18", colors["accent"]),
            "bottom": ("single", "12", colors["accent"]),
            "insideH": ("single", "4", colors["rule"]),
            "left": ("single", "4", colors["rule"]) if vertical else ("none", "0", "auto"),
            "right": ("single", "4", colors["rule"]) if vertical else ("none", "0", "auto"),
            "insideV": ("single", "4", colors["rule"]) if vertical else ("none", "0", "auto"),
        }
        for edge, (val, sz, color) in spec.items():
            el = self.OxmlElement(f"w:{edge}")
            el.set(self.qn("w:val"), val)
            el.set(self.qn("w:sz"), sz)
            el.set(self.qn("w:color"), color)
            borders.append(el)
        table._tbl.tblPr.append(borders)

    def write_title(self, title: str) -> None:
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        para = self.doc.add_paragraph()
        para.alignment = self.ALIGN.CENTER
        para.paragraph_format.space_after = self.Pt(16)
        self.font(para.add_run(title), name=f["fonts"]["title"],
                  size=sizes["title"], bold=True, color=colors["accent"])
        if f["rules"].get("title_rule"):
            self.bottom_rule(para, colors["accent"], 18)

    def write_section(self, index: int, name: str, status: str) -> None:
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        numeral = ROMAN[index] if index < len(ROMAN) else str(index + 1)
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = self.Pt(16)
        para.paragraph_format.space_after = self.Pt(6)
        self.font(para.add_run(f"{numeral}. {name}"), name=f["fonts"]["title"],
                  size=sizes["section"], bold=True, color=colors["accent"])
        if status and status != "확정":
            self.font(para.add_run(f"  ({status})"), name=f["fonts"]["body"],
                      size=sizes["note"], color=colors["worsen"])
        if f["rules"].get("section_rule"):
            self.bottom_rule(para, colors["accent"], 12)

    def write_block(self, label: str) -> None:
        """A `< >` block heading, outside the numbered chain."""
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = self.Pt(16)
        para.paragraph_format.space_after = self.Pt(5)
        self.font(para.add_run(f"< {label} >"), name=f["fonts"]["title"],
                  size=sizes["section"], bold=True, color=colors["accent"])

    def write_appendix(self, title: str, blocks: list) -> None:
        """Supporting material, after the argument and set smaller."""
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        head = self.doc.add_paragraph()
        head.paragraph_format.space_before = self.Pt(14)
        head.paragraph_format.space_after = self.Pt(4)
        head.paragraph_format.line_spacing = f["line_spacing"]["base"]
        self.font(head.add_run(title), name=f["fonts"]["title"],
                  size=sizes["appendix_title"], bold=True, color=colors["accent"])
        for level, payload in blocks:
            if level == "table":
                self.write_table(payload, size=sizes["appendix"])
            else:
                self.write_line(level, payload, size=sizes["appendix"])

    def paint_ground(self) -> None:
        """Fill the page. The reference sits on a tinted ground, not on white."""
        ground = self.form["colors"].get("ground", "FFFFFF")
        if ground.upper() == "FFFFFF":
            return
        bg = self.OxmlElement("w:background")
        bg.set(self.qn("w:color"), ground)
        self.doc.element.insert(0, bg)
        # CT_Settings orders displayBackgroundShape straight after w:zoom.
        settings = self.doc.settings.element
        show = self.OxmlElement("w:displayBackgroundShape")
        zoom = settings.find(self.qn("w:zoom"))
        if zoom is not None:
            zoom.addnext(show)
        else:
            settings.insert(0, show)

    def _right_tab(self, para) -> None:
        section = self.doc.sections[0]
        width = section.page_width - section.left_margin - section.right_margin
        para.paragraph_format.tab_stops.add_tab_stop(width, self.TAB_RIGHT)

    def write_chrome(self, dept: str, doc_name: str, nav: str = "") -> None:
        f, colors, sizes = self.form, self.form["colors"], self.form["sizes"]
        chrome = f.get("chrome", {})
        section = self.doc.sections[0]
        if chrome.get("header"):
            para = section.header.paragraphs[0]
            para.paragraph_format.space_after = self.Pt(2)
            self._right_tab(para)
            self.font(para.add_run(f"{dept}"), name=f["fonts"]["body"],
                      size=sizes["chrome"], bold=True, color=colors["accent"])
            if doc_name:
                self.font(para.add_run(f"  |  {doc_name}"), name=f["fonts"]["body"],
                          size=sizes["chrome"], color=colors["text"])
            if nav:
                self.font(para.add_run(f"\t{nav}"), name=f["fonts"]["body"],
                          size=sizes["chrome"], color=colors["rule"])
            self.bottom_rule(para, colors["rule"], 6)
        if chrome.get("footer"):
            para = section.footer.paragraphs[0]
            para.alignment = self.ALIGN.RIGHT
            self.top_rule(para, colors["rule"], 6)
            self.font(para.add_run(f"{dept}   |   "), name=f["fonts"]["body"],
                      size=sizes["chrome"], color=colors["accent"])
            self.font(para.add_run("- "), name=f["fonts"]["body"],
                      size=sizes["chrome"], color=colors["text"])
            self._page_field(para)
            self.font(para.add_run(" -"), name=f["fonts"]["body"],
                      size=sizes["chrome"], color=colors["text"])

    def _page_field(self, para) -> None:
        run = self.font(para.add_run(), name=self.form["fonts"]["body"],
                        size=self.form["sizes"]["chrome"],
                        color=self.form["colors"]["text"])
        begin = self.OxmlElement("w:fldChar")
        begin.set(self.qn("w:fldCharType"), "begin")
        instr = self.OxmlElement("w:instrText")
        instr.set(self.qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = self.OxmlElement("w:fldChar")
        end.set(self.qn("w:fldCharType"), "end")
        for el in (begin, instr, end):
            run._element.append(el)

    def _seal_auto_space(self) -> None:
        """Repeat the no-padding setting on every paragraph, table cells included.

        One pass at the end rather than at each of the places a paragraph is
        made: a new kind of paragraph added later would otherwise quietly opt
        out of it, and the defect only shows up in a rendered document.
        """
        for para in self.doc.paragraphs:
            self.no_auto_space_para(para)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.no_auto_space_para(para)

    def save(self, path: Path) -> None:
        self._seal_auto_space()
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
